import json
import logging

import function_app as fa
from function_app import app, ToolProperty, _init_word_libs
from storage_utils import (
    _download_blob_to_temp,
    _generate_blob_sas_url,
    _init_storage,
    _upload_file_to_blob,
)


def _ensure_cell_shading(cell, fill_color: str, pattern: str = "clear"):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), pattern)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(fill_color))
    tc_pr.append(shd)


def _set_cell_padding(cell, top=None, bottom=None, left=None, right=None, unit=None):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    def to_twips(val):
        try:
            return str(int(float(val) * 20))
        except Exception:
            return None

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in ("top", top), ("bottom", bottom), ("left", left), ("right", right):
        if value is None:
            continue
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        tw = to_twips(value)
        if tw is not None:
            el.set(qn("w:w"), tw)
            el.set(qn("w:type"), "dxa")


def _set_table_width(table, width_value: float, width_type: str = "points"):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    def to_twips(val):
        try:
            return str(int(float(val) * 20))
        except Exception:
            return None

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    if (width_type or "").lower() == "percent":
        try:
            pct = str(int(float(width_value) * 50))
            tbl_w.set(qn("w:w"), pct)
            tbl_w.set(qn("w:type"), "pct")
        except Exception as exc:
            logging.warning(
                "Failed to set table width to %s (file: unknown): %s",
                width_value,
                exc,
            )
    else:
        tw = to_twips(width_value)
        if tw is not None:
            tbl_w.set(qn("w:w"), tw)
            tbl_w.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_value: float, width_type: str = "points"):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    def to_twips(val):
        try:
            return str(int(float(val) * 20))
        except Exception:
            return None

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    if (width_type or "").lower() == "percent":
        try:
            pct = str(int(float(width_value) * 50))
            tc_w.set(qn("w:w"), pct)
            tc_w.set(qn("w:type"), "pct")
        except Exception as exc:
            logging.warning(
                "Failed to set cell width to %s (file: unknown): %s",
                width_value,
                exc,
            )
    else:
        tw = to_twips(width_value)
        if tw is not None:
            tc_w.set(qn("w:w"), tw)
            tc_w.set(qn("w:type"), "dxa")


def _set_table_layout_autofit(table):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")
    try:
        table.autofit = True
    except Exception as exc:
        logging.warning(
            "Failed to enable table autofit (file: unknown): %s",
            exc,
        )


word_tool_props_format_table = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("has_header_row", "boolean",
                 "Mark first row as header (visual only)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_format_table",
    description="Basic table formatting (header flag only).",
    toolProperties=word_tool_props_format_table,
)
def word_format_table(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    table_index = int(args.get("table_index", -1))
    has_header_row = bool(args.get("has_header_row", False))
    if not filename or table_index < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if table_index >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[table_index]
    if has_header_row and len(table.rows) > 0:
        try:
            for p in table.rows[0].cells[0].paragraphs:
                logging.warning(
                    "No operation for header paragraph in file %s (user %s)",
                    blob_name,
                    user_id,
                )
            for cell in table.rows[0].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        except Exception as exc:
            logging.warning(
                "Failed to format header row for %s (user %s): %s",
                blob_name,
                user_id,
                exc,
            )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_cell_shading = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("fill_color", "string", "Hex color without #").to_dict(),
    ToolProperty("pattern", "string", "Pattern (default clear)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_cell_shading",
    description="Set shading on a given table cell.",
    toolProperties=word_tool_props_set_table_cell_shading,
)
def word_set_table_cell_shading(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r_idx = int(args.get("row_index", -1))
    c_idx = int(args.get("col_index", -1))
    fill_color = args.get("fill_color")
    pattern = args.get("pattern", "clear")
    if not filename or t_idx < 0 or r_idx < 0 or c_idx < 0 or not fill_color:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if r_idx >= len(table.rows) or c_idx >= len(table.columns):
        return "cell index out of range"
    _ensure_cell_shading(table.cell(r_idx, c_idx), fill_color, pattern)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_apply_table_alternating_rows = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("color1", "string", "Hex color 1 (default FFFFFF)").to_dict(),
    ToolProperty("color2", "string", "Hex color 2 (default F2F2F2)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_apply_table_alternating_rows",
    description="Apply alternating row colors.",
    toolProperties=word_tool_props_apply_table_alternating_rows,
)
def word_apply_table_alternating_rows(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    color1 = args.get("color1", "FFFFFF")
    color2 = args.get("color2", "F2F2F2")
    if not filename or t_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    for r_idx, row in enumerate(table.rows):
        fill = color1 if r_idx % 2 == 0 else color2
        for c_idx, _ in enumerate(table.columns):
            _ensure_cell_shading(table.cell(r_idx, c_idx), fill, "clear")
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_highlight_table_header = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("header_color", "string",
                 "Header color (default 4472C4)").to_dict(),
    ToolProperty("text_color", "string",
                 "Header text color (default FFFFFF)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_highlight_table_header",
    description="Highlight first row as header with colors.",
    toolProperties=word_tool_props_highlight_table_header,
)
def word_highlight_table_header(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    header_color = args.get("header_color", "4472C4")
    text_color = args.get("text_color", "FFFFFF")
    if not filename or t_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables) or len(doc.tables[t_idx].rows) == 0:
        return "table_index out of range"
    table = doc.tables[t_idx]
    for cell in table.rows[0].cells:
        _ensure_cell_shading(cell, header_color, "clear")
        try:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = fa._docx.shared.RGBColor.from_string(text_color)
        except Exception:
            pass
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_merge_cells = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("start_row", "number", "Start row index").to_dict(),
    ToolProperty("start_col", "number", "Start column index").to_dict(),
    ToolProperty("end_row", "number", "End row index").to_dict(),
    ToolProperty("end_col", "number", "End column index").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_merge_table_cells",
    description="Merge a rectangular range of cells.",
    toolProperties=word_tool_props_merge_cells,
)
def word_merge_table_cells(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    sr = int(args.get("start_row", -1))
    sc = int(args.get("start_col", -1))
    er = int(args.get("end_row", -1))
    ec = int(args.get("end_col", -1))
    if not filename or t_idx < 0 or sr < 0 or sc < 0 or er < 0 or ec < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    try:
        table.cell(sr, sc).merge(table.cell(er, ec))
    except Exception as exc:
        return f"merge failed: {exc}"
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_merge_cells_horizontal = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index").to_dict(),
    ToolProperty("start_col", "number", "Start column index").to_dict(),
    ToolProperty("end_col", "number", "End column index").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_merge_table_cells_horizontal",
    description="Merge horizontally in a single row.",
    toolProperties=word_tool_props_merge_cells_horizontal,
)
def word_merge_table_cells_horizontal(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r = int(args.get("row_index", -1))
    sc = int(args.get("start_col", -1))
    ec = int(args.get("end_col", -1))
    if not filename or t_idx < 0 or r < 0 or sc < 0 or ec < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    try:
        table.cell(r, sc).merge(table.cell(r, ec))
    except Exception as exc:
        return f"merge failed: {exc}"
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_merge_cells_vertical = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index").to_dict(),
    ToolProperty("start_row", "number", "Start row index").to_dict(),
    ToolProperty("end_row", "number", "End row index").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_merge_table_cells_vertical",
    description="Merge vertically in a single column.",
    toolProperties=word_tool_props_merge_cells_vertical,
)
def word_merge_table_cells_vertical(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    c = int(args.get("col_index", -1))
    sr = int(args.get("start_row", -1))
    er = int(args.get("end_row", -1))
    if not filename or t_idx < 0 or c < 0 or sr < 0 or er < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    try:
        table.cell(sr, c).merge(table.cell(er, c))
    except Exception as exc:
        return f"merge failed: {exc}"
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_cell_alignment = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("horizontal", "string",
                 "left|center|right|justify").to_dict(),
    ToolProperty("vertical", "string", "top|center|bottom").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_cell_alignment",
    description="Set horizontal/vertical alignment for a cell.",
    toolProperties=word_tool_props_set_table_cell_alignment,
)
def word_set_table_cell_alignment(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r = int(args.get("row_index", -1))
    c = int(args.get("col_index", -1))
    horizontal = (args.get("horizontal") or "").lower()
    vertical = (args.get("vertical") or "").lower()
    if not filename or t_idx < 0 or r < 0 or c < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if r >= len(table.rows) or c >= len(table.columns):
        return "cell index out of range"
    cell = table.cell(r, c)
    try:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if horizontal:
            mapping = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            align = mapping.get(horizontal)
            if align is not None:
                for p in cell.paragraphs:
                    p.alignment = align
        if vertical:
            vm = {
                "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
                "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
            }
            v = vm.get(vertical)
            if v is not None:
                cell.vertical_alignment = v
    except Exception as exc:
        logging.warning(
            "Failed to set cell alignment for %s (user %s): %s",
            blob_name,
            user_id,
            exc,
        )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_alignment_all = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("horizontal", "string",
                 "left|center|right|justify").to_dict(),
    ToolProperty("vertical", "string", "top|center|bottom").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_alignment_all",
    description="Apply alignment to all cells in a table.",
    toolProperties=word_tool_props_set_table_alignment_all,
)
def word_set_table_alignment_all(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    horizontal = (args.get("horizontal") or "").lower()
    vertical = (args.get("vertical") or "").lower()
    if not filename or t_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    try:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        hmap = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        vmap = {
            "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
            "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        }
        halign = hmap.get(horizontal)
        valign = vmap.get(vertical)
        for row in table.rows:
            for cell in row.cells:
                if halign is not None:
                    for p in cell.paragraphs:
                        p.alignment = halign
                if valign is not None:
                    cell.vertical_alignment = valign
    except Exception as exc:
        logging.warning(
            "Failed to set table alignment for %s (user %s): %s",
            blob_name,
            user_id,
            exc,
        )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_format_table_cell_text = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("text_content", "string", "Replace cell text").to_dict(),
    ToolProperty("bold", "boolean", "Bold").to_dict(),
    ToolProperty("italic", "boolean", "Italic").to_dict(),
    ToolProperty("underline", "boolean", "Underline").to_dict(),
    ToolProperty("color", "string", "Hex color without #").to_dict(),
    ToolProperty("font_size", "number", "Font size in points").to_dict(),
    ToolProperty("font_name", "string", "Font family name").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_format_table_cell_text",
    description="Format text in a specific table cell.",
    toolProperties=word_tool_props_format_table_cell_text,
)
def word_format_table_cell_text(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r = int(args.get("row_index", -1))
    c = int(args.get("col_index", -1))
    text = args.get("text_content")
    bold = args.get("bold")
    italic = args.get("italic")
    underline = args.get("underline")
    color = args.get("color")
    font_size = args.get("font_size")
    font_name = args.get("font_name")
    if not filename or t_idx < 0 or r < 0 or c < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if r >= len(table.rows) or c >= len(table.columns):
        return "cell index out of range"
    cell = table.cell(r, c)
    if text is not None:
        cell.text = str(text)
    try:
        for p in cell.paragraphs:
            for run in p.runs:
                if bold is not None:
                    run.bold = bool(bold)
                if italic is not None:
                    run.italic = bool(italic)
                if underline is not None:
                    run.underline = bool(underline)
                if color:
                    try:
                        run.font.color.rgb = fa._docx.shared.RGBColor.from_string(str(color))
                    except Exception:
                        pass
                if font_size is not None:
                    try:
                        run.font.size = fa._docx.shared.Pt(float(font_size))
                    except Exception as exc:
                        logging.warning(
                            "Failed to set font size for %s (user %s): %s",
                            blob_name,
                            user_id,
                            exc,
                        )
                if font_name:
                    run.font.name = str(font_name)
    except Exception as exc:
        logging.error(
            "Failed to format table cell text for %s (user %s): %s",
            blob_name,
            user_id,
            exc,
        )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_cell_padding = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("top", "number", "Top padding (points)").to_dict(),
    ToolProperty("bottom", "number", "Bottom padding (points)").to_dict(),
    ToolProperty("left", "number", "Left padding (points)").to_dict(),
    ToolProperty("right", "number", "Right padding (points)").to_dict(),
    ToolProperty("unit", "string",
                 "Padding unit in points (only 'points' supported)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_cell_padding",
    description="Set padding for a specific table cell.",
    toolProperties=word_tool_props_set_table_cell_padding,
)
def word_set_table_cell_padding(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r = int(args.get("row_index", -1))
    c = int(args.get("col_index", -1))
    top = args.get("top")
    bottom = args.get("bottom")
    left = args.get("left")
    right = args.get("right")
    unit = args.get("unit") or "points"
    if not filename or t_idx < 0 or r < 0 or c < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if r >= len(table.rows) or c >= len(table.columns):
        return "cell index out of range"
    _set_cell_padding(table.cell(r, c), top=top, bottom=bottom,
                      left=left, right=right, unit=unit)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_column_width = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("width", "number",
                 "Width value (points or percent)").to_dict(),
    ToolProperty("width_type", "string",
                 "points|percent (default points)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_column_width",
    description="Set width for a specific column.",
    toolProperties=word_tool_props_set_table_column_width,
)
def word_set_table_column_width(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    c = int(args.get("col_index", -1))
    width = args.get("width")
    width_type = args.get("width_type") or "points"
    if not filename or t_idx < 0 or c < 0 or width is None:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if c >= len(table.columns):
        return "col_index out of range"
    for row in table.rows:
        _set_cell_width(row.cells[c], width, width_type)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_column_widths = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("widths", "array", "Array of widths by column (points unless width_type=percent)",
                 item_type="number").to_dict(),
    ToolProperty("width_type", "string",
                 "points|percent (default points)").to_dict(),
])


word_tool_props_set_table_width = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("width", "number", "Table width value").to_dict(),
    ToolProperty("width_type", "string",
                 "points|percent (default points)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_width",
    description="Set overall table width.",
    toolProperties=word_tool_props_set_table_width,
)
def word_set_table_width(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    width = args.get("width")
    width_type = args.get("width_type") or "points"
    if not filename or t_idx < 0 or width is None:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    _set_table_width(table, width, width_type)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_auto_fit_table_columns = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_auto_fit_table_columns",
    description="Enable table auto-fit layout.",
    toolProperties=word_tool_props_auto_fit_table_columns,
)
def word_auto_fit_table_columns(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    if not filename or t_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    _set_table_layout_autofit(table)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_add_table = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("rows", "number", "Number of rows").to_dict(),
        ToolProperty("cols", "number", "Number of columns").to_dict(),
        ToolProperty(
            "data",
            "string",
            "Optional JSON 2D array of cell texts, e.g. [[\"A\",\"B\"],[\"C\",\"D\"]]",
        ).to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_add_table",
    description="Insert a table; optionally populate with provided data.",
    toolProperties=word_tool_props_add_table,
)
def word_add_table(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    rows = int(args.get("rows", 0))
    cols = int(args.get("cols", 0))
    data = args.get("data")
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except Exception:
            data = None
    if not filename or rows <= 0 or cols <= 0:
        return "Missing filename or invalid rows/cols"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = fa._docx.Document(local_path)
    table = doc.add_table(rows=rows, cols=cols)
    if isinstance(data, list):
        for r_idx in range(min(rows, len(data))):
            row_data = data[r_idx] if isinstance(data[r_idx], list) else []
            for c_idx in range(min(cols, len(row_data))):
                table.cell(r_idx, c_idx).text = str(row_data[c_idx])
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})












