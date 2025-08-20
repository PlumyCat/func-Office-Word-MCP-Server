import json
import logging
import os
import uuid
from datetime import datetime, timedelta, timezone

import azure.functions as func

app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)
"""Word MCP Server tools (Azure Functions v2 decorators)."""


class ToolProperty:
    def __init__(self, property_name: str, property_type: str, description: str, item_type: str | None = None):
        self.propertyName = property_name
        self.propertyType = property_type
        self.description = description
        self.itemType = item_type

    def to_dict(self):
        d = {
            "propertyName": self.propertyName,
            "propertyType": self.propertyType,
            "description": self.description,
        }
        if self.propertyType == "array" and self.itemType:
            # JSON Schema-like hint for items type
            d["items"] = {"type": self.itemType}
        return d


# (Removed snippet demo tool properties)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="hello_mcp",
    description="Hello world.",
    toolProperties="[]",
)
def hello_mcp(context) -> None:
    """
    A simple function that returns a greeting message.

    Args:
        context: The trigger context (not used in this function).

    Returns:
        str: A greeting message.
    """
    return "Hello I am MCPTool!"




# Health route (anonymous)
@app.route(route="ping", methods=["GET"], auth_level=func.AuthLevel.ANONYMOUS)
def ping(req: func.HttpRequest) -> func.HttpResponse:
    return func.HttpResponse(
        body=json.dumps({"status": "ok"}),
        mimetype="application/json",
        status_code=200,
    )


# ---- Word MCP minimal tools (lazy imports) ----
_docx = None
_blob_service_client = None
_blob_container_name = None
_ORG_TEMPLATES_PREFIX = os.environ.get(
    "ORG_TEMPLATES_PREFIX", "shared/templates/")
_graph_session = None
_graph_token_expires = None

# Template upload validation defaults
_ALLOWED_TEMPLATE_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template",  # .dotx
}
_ALLOWED_TEMPLATE_EXTENSIONS = {".docx", ".dotx"}


def _init_word_libs():
    global _docx
    if _docx is not None:
        return
    try:
        import docx as _docx_mod  # python-docx
    except Exception as exc:
        logging.error("Failed to import python-docx: %s", exc)
        raise
    _docx = _docx_mod


def _init_storage():
    global _blob_service_client, _blob_container_name
    if _blob_service_client is not None:
        return
    try:
        from azure.storage.blob import BlobServiceClient
        try:
            from azure.core.pipeline.transport import RequestsTransport
        except Exception:
            RequestsTransport = None
    except Exception as exc:
        logging.error("Failed to import azure-storage-blob: %s", exc)
        raise
    connection_string = os.environ.get(
        "WORD_STORAGE_CONNECTION_STRING") or os.environ.get("AzureWebJobsStorage")
    if not connection_string:
        raise RuntimeError("AzureWebJobsStorage is not configured")
    _blob_container_name = os.environ.get("WORD_DOCS_CONTAINER", "stword")
    if RequestsTransport is not None:
        try:
            conn_timeout = float(os.environ.get("BLOB_CONN_TIMEOUT", "60"))
            read_timeout = float(os.environ.get("BLOB_READ_TIMEOUT", "300"))
            transport = RequestsTransport(
                connection_timeout=conn_timeout, read_timeout=read_timeout)
            _blob_service_client = BlobServiceClient.from_connection_string(
                connection_string, transport=transport)
        except Exception:
            _blob_service_client = BlobServiceClient.from_connection_string(
                connection_string)
    else:
        _blob_service_client = BlobServiceClient.from_connection_string(
            connection_string)
    try:
        _blob_service_client.create_container(_blob_container_name)
    except Exception as exc:
        # Likely already exists or insufficient permissions; continue
        logging.warning(
            "Could not create container %s: %s", _blob_container_name, exc
        )


def _get_blob_client(blob_name: str):
    _init_storage()
    container_client = _blob_service_client.get_container_client(
        _blob_container_name)
    return container_client.get_blob_client(blob_name)


def _blob_exists(blob_name: str) -> bool:
    try:
        client = _get_blob_client(blob_name)
        try:
            return bool(client.exists())
        except Exception:
            # Fallback for older SDKs
            client.get_blob_properties()
            return True
    except Exception:
        return False


def _with_retries(operation, *, max_attempts: int = 3, base_delay_seconds: float = 0.5):
    """Run an idempotent operation with basic exponential backoff."""
    attempt = 0
    last_exc = None
    while attempt < max_attempts:
        try:
            return operation()
        except Exception as exc:
            last_exc = exc
            attempt += 1
            if attempt >= max_attempts:
                break
            sleep_for = base_delay_seconds * (2 ** (attempt - 1))
            try:
                import time
                time.sleep(sleep_for)
            except Exception as exc:
                logging.warning(
                    "Retry sleep failed for operation %s (user: unknown, file: unknown): %s",
                    getattr(operation, "__name__", repr(operation)),
                    exc,
                )
    if last_exc:
        raise last_exc
    return None


def _ensure_user_paths(user_id: str) -> list[str]:
    """Ensure the user prefix and required sub-prefixes exist by creating placeholder blobs."""
    _init_storage()
    created: list[str] = []
    placeholders = [
        f"{user_id}/.keep",
        f"{user_id}/image_blob/.keep",
        f"{user_id}/templates/.keep",
    ]
    for ph in placeholders:
        try:
            client = _get_blob_client(ph)

            def _op():
                return client.upload_blob(b"", overwrite=False)
            _with_retries(_op)
            created.append(ph)
        except Exception as exc:
            # Exists or not allowed; ignore
            logging.warning("Failed to create placeholder %s: %s", ph, exc)
    return created


def _parse_storage_connection_string(connection_string: str) -> dict:
    parts = connection_string.split(";")
    kv = {}
    for part in parts:
        if not part:
            continue
        if "=" not in part:
            continue
        k, v = part.split("=", 1)
        kv[k.strip().lower()] = v.strip()
    return kv


def _generate_blob_sas_url(blob_name: str, permissions: str = "r") -> dict:
    _init_storage()
    try:
        from azure.storage.blob import generate_blob_sas, BlobSasPermissions
    except Exception as exc:
        logging.error(
            "Failed to import azure-storage-blob SAS helpers: %s", exc)
        raise
    conn = os.environ.get("WORD_STORAGE_CONNECTION_STRING") or os.environ.get(
        "AzureWebJobsStorage")
    if not conn:
        raise RuntimeError(
            "No storage connection string available for SAS generation")
    kv = _parse_storage_connection_string(conn)
    account_name = kv.get("accountname")
    account_key = kv.get("accountkey")
    if not account_name or not account_key:
        raise RuntimeError(
            "Storage connection string missing AccountName/AccountKey for SAS generation")
    ttl_seconds_str = os.environ.get("WORD_BLOB_TTL_SECONDS", "3600")
    try:
        ttl_seconds = int(ttl_seconds_str)
    except Exception:
        ttl_seconds = 3600
    expiry = datetime.now(timezone.utc) + timedelta(seconds=ttl_seconds)
    sas_token = generate_blob_sas(
        account_name=account_name,
        container_name=_blob_container_name,
        blob_name=blob_name,
        account_key=account_key,
        permission=BlobSasPermissions.from_string(permissions),
        expiry=expiry,
    )
    blob_url = _get_blob_client(blob_name).url
    return {"url": f"{blob_url}?{sas_token}", "expiresUtc": expiry.isoformat()}


def _word_doc_path(filename: str) -> str:
    safe_name = os.path.basename(filename)
    return os.path.join("/tmp", safe_name)


def _download_blob_to_temp(blob_name: str) -> str:
    blob_client = _get_blob_client(blob_name)
    temp_path = _word_doc_path(blob_name)
    os.makedirs(os.path.dirname(temp_path), exist_ok=True)
    with open(temp_path, "wb") as f:
        data = blob_client.download_blob()
        f.write(data.readall())
    return temp_path


def _upload_file_to_blob(local_path: str, blob_name: str) -> None:
    blob_client = _get_blob_client(blob_name)
    with open(local_path, "rb") as f:
        blob_client.upload_blob(f, overwrite=True, max_concurrency=1)
    _apply_blob_ttl(blob_client)


def _apply_blob_ttl(blob_client) -> None:
    ttl_seconds_str = os.environ.get("WORD_BLOB_TTL_SECONDS", "3600")
    try:
        ttl_seconds = int(ttl_seconds_str)
    except Exception:
        ttl_seconds = 3600
    if ttl_seconds <= 0:
        return
    expires_on = datetime.now(timezone.utc) + timedelta(seconds=ttl_seconds)
    # Try SDK-level expiry; fall back to metadata if unavailable
    try:
        # Some SDK versions use set_blob_expiry(expiry_time=...); others use expires_on
        try:
            blob_client.set_blob_expiry(expires_on=expires_on)
            return
        except TypeError:
            blob_client.set_blob_expiry(expiry_time=expires_on)
            return
    except Exception as exc:
        logging.warning(
            "Failed to set expiry for blob %s: %s",
            getattr(blob_client, "blob_name", "unknown"),
            exc,
        )
    try:
        blob_client.set_blob_metadata({"expiry_utc": expires_on.isoformat()})
    except Exception as exc:
        logging.warning(
            "Failed to set metadata expiry for blob %s: %s",
            getattr(blob_client, "blob_name", "unknown"),
            exc,
        )


def _upload_bytes_to_blob(blob_client, data: bytes, content_settings=None) -> None:
    def _op():
        return blob_client.upload_blob(
            data,
            overwrite=True,
            content_settings=content_settings,
            max_concurrency=1,
        )
    _with_retries(_op)


def _init_graph_session():
    global _graph_session, _graph_token_expires
    if _graph_session is not None and _graph_token_expires:
        try:
            from datetime import datetime, timezone
            if datetime.now(timezone.utc) < _graph_token_expires:
                return
        except Exception as exc:
            logging.warning(
                "Graph session expiry check failed (user: N/A, file: N/A): %s",
                exc,
            )
    try:
        import requests
        import msal
    except Exception as exc:
        logging.error("Failed to import Graph deps: %s", exc)
        raise
    tenant_id = os.environ.get("TENANT_ID")
    client_id = os.environ.get("CLIENT_ID")
    client_secret = os.environ.get("CLIENT_SECRET")
    if not tenant_id or not client_id or not client_secret:
        raise RuntimeError(
            "Graph app credentials missing (TENANT_ID, CLIENT_ID, CLIENT_SECRET)")
    authority = f"https://login.microsoftonline.com/{tenant_id}"
    app = msal.ConfidentialClientApplication(
        client_id=client_id, client_credential=client_secret, authority=authority)
    scope = ["https://graph.microsoft.com/.default"]
    result = app.acquire_token_silent(scopes=scope, account=None)
    if not result:
        result = app.acquire_token_for_client(scopes=scope)
    if "access_token" not in result:
        raise RuntimeError(f"Graph token acquisition failed: {result}")
    token = result["access_token"]
    try:
        from datetime import datetime, timezone, timedelta
        _graph_token_expires = datetime.now(
            timezone.utc) + timedelta(seconds=int(result.get("expires_in", 300)))
    except Exception:
        _graph_token_expires = None
    sess = requests.Session()
    sess.headers.update(
        {"Authorization": f"Bearer {token}", "Accept": "application/json"})
    _graph_session = sess


def _graph_request(method: str, url: str, **kwargs):
    _init_graph_session()
    import requests
    sess = _graph_session
    resp = sess.request(method, url, timeout=30, **kwargs)
    if resp.status_code == 401:
        # Refresh token once
        _init_graph_session()
        sess = _graph_session
        resp = sess.request(method, url, timeout=30, **kwargs)
    if resp.status_code >= 400:
        raise RuntimeError(f"Graph HTTP {resp.status_code}: {resp.text[:500]}")
    return resp


def _graph_upload_to_drive(drive_id: str, name: str, content: bytes) -> dict:
    # Small files upload (<= 4MB) via simple upload; for larger, create upload session
    if len(content) <= 4 * 1024 * 1024:
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{name}:/content"
        resp = _graph_request("PUT", url, data=content, headers={
                              "Content-Type": "application/octet-stream"})
        return resp.json()
    # Upload session for big files
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{name}:/createUploadSession"
    session = _graph_request("POST", url, json={
                             "item": {"@microsoft.graph.conflictBehavior": "replace"}}).json()
    upload_url = session.get("uploadUrl")
    if not upload_url:
        raise RuntimeError("No uploadUrl in session")
    import math
    import requests
    chunk = 5 * 1024 * 1024
    size = len(content)
    start = 0
    while start < size:
        end = min(start + chunk, size) - 1
        headers = {
            "Content-Length": str(end - start + 1),
            "Content-Range": f"bytes {start}-{end}/{size}",
        }
        r = requests.put(upload_url, headers=headers,
                         data=content[start:end+1], timeout=60)
        if r.status_code not in (200, 201, 202):
            raise RuntimeError(
                f"Upload chunk failed: {r.status_code} {r.text[:200]}")
        start = end + 1
    # Get item
    # Final 201/200 may already include item JSON; fetch to be safe
    get_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{name}"
    return _graph_request("GET", get_url).json()


def _graph_download_pdf_content(drive_id: str, item_id: str) -> bytes:
    # Always target the specific drive we uploaded to
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/content?format=pdf"
    resp = _graph_request("GET", url, stream=True)
    return resp.content


def _graph_delete_item(drive_id: str, item_id: str):
    # Delete explicitly from the drive we used
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}"
    _graph_request("DELETE", url)


def _resolve_sharepoint_drive() -> str:
    drive_id = os.environ.get("SP_DRIVE_ID")
    if drive_id:
        return drive_id
    site_id = os.environ.get("SP_SITE_ID")
    if not site_id:
        raise RuntimeError("SP_DRIVE_ID or SP_SITE_ID must be configured")
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive"
    data = _graph_request("GET", url).json()
    return data.get("id")


# Define tool properties JSON for Word tools
word_tool_props_create = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Filename for the new .docx (optional; UUID if omitted)").to_dict(),
        ToolProperty("title", "string", "Optional document title").to_dict(),
        ToolProperty("author", "string", "Optional author").to_dict(),
        ToolProperty("template_blob", "string",
                     "Optional template blob path, e.g. templates/base.dotx").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_create_document",
    description="Create a new Word document and save it in /tmp.",
    toolProperties=word_tool_props_create,
)
def word_create_document(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename") or f"{uuid.uuid4()}.docx"
    title = args.get("title")
    author = args.get("author")
    template_blob = args.get("template_blob")
    qualified_blob_name = f"{user_id}/{filename}" if user_id else filename
    if template_blob:
        # Server-side copy from template; then optionally set core properties
        try:
            src_blob = _get_blob_client(template_blob)
            dest_blob = _get_blob_client(qualified_blob_name)
            dest_blob.start_copy_from_url(src_blob.url)
            if title or author:
                try:
                    local_path = _download_blob_to_temp(qualified_blob_name)
                    doc = _docx.Document(local_path)
                    if title:
                        try:
                            doc.core_properties.title = title
                        except Exception as exc:
                            logging.warning(
                                "Failed to set title for %s: %s",
                                qualified_blob_name,
                                exc,
                            )
                    if author:
                        try:
                            doc.core_properties.author = author
                        except Exception as exc:
                            logging.warning(
                                "Failed to set author for %s: %s",
                                qualified_blob_name,
                                exc,
                            )
                    doc.save(local_path)
                    _upload_file_to_blob(local_path, qualified_blob_name)
                except Exception as exc:
                    # Ignore property set failures on template-produced docs
                    logging.warning(
                        "Failed to apply core properties to %s: %s",
                        qualified_blob_name,
                        exc,
                    )
        except Exception:
            # Fallback: create a blank document
            local_path = _word_doc_path(filename)
            doc = _docx.Document()
            if title:
                try:
                    doc.core_properties.title = title
                except Exception as exc:
                    logging.warning(
                        "Failed to set title for %s: %s",
                        qualified_blob_name,
                        exc,
                    )
            if author:
                try:
                    doc.core_properties.author = author
                except Exception as exc:
                    logging.warning(
                        "Failed to set author for %s: %s",
                        qualified_blob_name,
                        exc,
                    )
            doc.save(local_path)
            _upload_file_to_blob(local_path, qualified_blob_name)
    else:
        local_path = _word_doc_path(filename)
        doc = _docx.Document()
        if title:
            try:
                doc.core_properties.title = title
            except Exception as exc:
                logging.warning(
                    "Failed to set title for %s: %s",
                    qualified_blob_name,
                    exc,
                )
        if author:
            try:
                doc.core_properties.author = author
            except Exception as exc:
                logging.warning(
                    "Failed to set author for %s: %s",
                    qualified_blob_name,
                    exc,
                )
        doc.save(local_path)
        _upload_file_to_blob(local_path, qualified_blob_name)
    # Return blob info and SAS URL for convenience
    sas = _generate_blob_sas_url(qualified_blob_name, permissions="r")
    response = {
        "container": _blob_container_name,
        "blob": qualified_blob_name,
        "sasUrl": sas.get("url"),
        "expiresUtc": sas.get("expiresUtc"),
    }
    return json.dumps(response)


# ---- New HTTP endpoints: user init, image upload/list, template upload/list ----


@app.route(route="users/init", methods=["POST"], auth_level=func.AuthLevel.FUNCTION)
def init_user(req: func.HttpRequest) -> func.HttpResponse:
    """Initialise l'espace utilisateur.

    Format attendu: POST /api/users/init  JSON {"user_id": "..."}
    (On tolère encore userId en entrée mais on renvoie user_id.)
    """
    user_id: str | None = None
    try:
        data = req.get_json()
        if isinstance(data, dict):
            user_id = data.get("user_id") or data.get("userId")
    except Exception:
        pass
    if not user_id:
        return func.HttpResponse("Missing user_id in JSON body", status_code=400)
    try:
        created = _ensure_user_paths(user_id)
        body = json.dumps({"user_id": user_id, "created": created})
        return func.HttpResponse(body=body, mimetype="application/json", status_code=200)
    except Exception as exc:
        logging.exception("init_user failed")
        return func.HttpResponse(f"init failed: {exc}", status_code=500)


def _read_form_file(req: func.HttpRequest, field_name: str = "file") -> tuple[bytes, str, str]:
    """Read raw bytes and filename from multipart/form-data or octet-stream."""
    user_hint = req.params.get("userId") or "unknown"
    # Prefer multipart (explicit parse; Azure Functions does not auto-parse files)
    try:
        ct_header = req.headers.get("content-type", "")
        if "multipart/form-data" in ct_header.lower():
            try:
                from requests_toolbelt.multipart import decoder as _mt_decoder  # lazy optional
                body = req.get_body() or b""
                dec = _mt_decoder.MultipartDecoder(body, ct_header)
                for part in dec.parts:
                    # Extract filename from Content-Disposition
                    filename = None
                    try:
                        cd = part.headers.get(
                            b"Content-Disposition", b"").decode("utf-8", "ignore")
                        for seg in cd.split(";"):
                            seg = seg.strip()
                            if seg.startswith("filename="):
                                filename = seg.split("=", 1)[1].strip('"')
                                break
                    except Exception:
                        filename = None
                    if not filename:
                        continue
                    content_type = None
                    try:
                        content_type = part.headers.get(
                            b"Content-Type", b"").decode("utf-8", "ignore") or None
                    except Exception:
                        content_type = None
                    data = part.content or b""
                    if data:
                        return data, filename, content_type or "application/octet-stream"
            except Exception as exc:
                logging.warning(
                    "Failed to parse multipart data for field %s (user %s): %s",
                    field_name,
                    user_hint,
                    exc,
                )
    except Exception as exc:
        logging.warning(
            "Failed to read multipart body for field %s (user %s): %s",
            field_name,
            user_hint,
            exc,
        )
    # Try framework-provided files (may not exist in Azure Functions)
    try:
        files = req.files or {}
        file = files.get(field_name)
        if file is not None:
            data = file.read()
            filename = getattr(file, "filename", None) or req.params.get(
                "fileName") or "upload.bin"
            content_type = getattr(file, "content_type", None) or req.headers.get(
                "content-type") or "application/octet-stream"
            return data, filename, content_type
    except Exception as exc:
        logging.warning(
            "Framework file parsing failed for field %s (user %s): %s",
            field_name,
            user_hint,
            exc,
        )
    # Fallback to raw body
    data = req.get_body() or b""
    filename = req.params.get("fileName") or "upload.bin"
    content_type = req.headers.get(
        "content-type") or "application/octet-stream"
    return data, filename, content_type


# Removed multi-file parsing helper; using single-file uploads only


def _sanitize_filename(name: str) -> str:
    base = os.path.basename(name or "")
    return base.replace("\\", "_").replace("/", "_")


def _guess_mime_type(filename: str, header_content_type: str | None) -> str:
    if header_content_type and header_content_type != "application/octet-stream":
        return header_content_type
    try:
        import mimetypes
        guessed, _ = mimetypes.guess_type(filename)
        if guessed:
            return guessed
    except Exception as exc:
        logging.warning(
            "Failed to guess MIME type for file %s: %s",
            filename,
            exc,
        )
    return "application/octet-stream"


def _make_content_settings(content_type: str):
    try:
        from azure.storage.blob import ContentSettings
        return ContentSettings(content_type=content_type)
    except Exception:
        return None


@app.route(route="users/images", methods=["POST"], auth_level=func.AuthLevel.FUNCTION)
def upload_image(req: func.HttpRequest) -> func.HttpResponse:
    # Exige user_id dans JSON. (Tolère userId.)
    user_id = None
    try:
        data = req.get_json()
        if isinstance(data, dict):
            user_id = data.get("user_id") or data.get("userId")
    except Exception:
        pass
    if not user_id:
        return func.HttpResponse("Missing user_id in JSON body", status_code=400)
    _init_storage()
    try:
        data, filename, content_type = _read_form_file(req, field_name="file")
        if not data:
            return func.HttpResponse("Empty body", status_code=400)
        try:
            max_mb = int(os.environ.get("IMAGE_MAX_MB", "10"))
        except Exception:
            max_mb = 10
        if len(data) > max_mb * 1024 * 1024:
            return func.HttpResponse(f"File too large (>{max_mb}MB): {filename}", status_code=413)
        safe_name = _sanitize_filename(req.params.get("fileName") or filename)
        blob_name = f"{user_id}/image_blob/{safe_name}"
        blob_client = _get_blob_client(blob_name)
        cs = _make_content_settings(_guess_mime_type(safe_name, content_type))
        _upload_bytes_to_blob(blob_client, data, content_settings=cs)
        _apply_blob_ttl(blob_client)
        props = blob_client.get_blob_properties()
        meta = {
            "blob": blob_name,
            "etag": getattr(props, "etag", None),
            "size": getattr(props, "size", None),
            "contentType": getattr(props, "content_settings", None).content_type if getattr(props, "content_settings", None) else content_type,
            "lastModified": getattr(props, "last_modified", None).isoformat() if getattr(props, "last_modified", None) else None,
        }
        sas = _generate_blob_sas_url(blob_name, permissions="r")
        meta["sasUrl"] = sas.get("url")
        meta["expiresUtc"] = sas.get("expiresUtc")
        return func.HttpResponse(body=json.dumps(meta), mimetype="application/json", status_code=201)
    except Exception as exc:
        logging.exception("upload_image failed")
        return func.HttpResponse(f"upload failed: {exc}", status_code=500)


@app.route(route="users/images", methods=["GET"], auth_level=func.AuthLevel.FUNCTION)
def list_images(req: func.HttpRequest) -> func.HttpResponse:
    # Exige user_id dans JSON body (GET) ou tolère userId; pas de query.
    user_id = None
    try:
        data = req.get_json()
        if isinstance(data, dict):
            user_id = data.get("user_id") or data.get("userId")
    except Exception:
        pass
    if not user_id:
        return func.HttpResponse("Missing user_id in JSON body", status_code=400)
    _init_storage()
    container_client = _blob_service_client.get_container_client(
        _blob_container_name)
    prefix = f"{user_id}/image_blob/"
    try:
        page_size_val: int = 50
        if isinstance(data, dict) and data.get("pageSize") is not None:
            try:
                page_size_val = int(data.get("pageSize"))
            except Exception:
                page_size_val = 50
        page_size = max(1, min(page_size_val, 200))
        token = req.params.get("continuationToken")
        pager = container_client.list_blobs(
            name_starts_with=prefix, results_per_page=page_size).by_page(continuation_token=token)
        items = []
        next_token = None
        for page in pager:
            for blob in page:
                if blob.name.endswith("/.keep"):
                    continue
                items.append({
                    "name": blob.name,
                    "size": getattr(blob, "size", None),
                    "contentType": getattr(getattr(blob, "content_settings", None), "content_type", None),
                    "lastModified": getattr(blob, "last_modified", None).isoformat() if getattr(blob, "last_modified", None) else None,
                })
            next_token = getattr(pager, "continuation_token", None)
            break
        payload = {"items": items}
        if next_token:
            payload["continuationToken"] = next_token
        return func.HttpResponse(body=json.dumps(payload), mimetype="application/json", status_code=200)
    except Exception as exc:
        logging.exception("list_images failed")
        return func.HttpResponse(f"list failed: {exc}", status_code=500)


@app.route(route="users/templates", methods=["POST"], auth_level=func.AuthLevel.FUNCTION)
def upload_template(req: func.HttpRequest) -> func.HttpResponse:
    user_id = None
    try:
        data = req.get_json()
        if isinstance(data, dict):
            user_id = data.get("user_id") or data.get("userId")
    except Exception:
        pass
    if not user_id:
        return func.HttpResponse("Missing user_id in JSON body", status_code=400)
    _init_storage()
    try:
        data, filename, content_type = _read_form_file(req, field_name="file")
        if not data:
            return func.HttpResponse("Empty body", status_code=400)
        try:
            max_mb = int(os.environ.get("TEMPLATE_MAX_MB", "10"))
        except Exception:
            max_mb = 10
        if len(data) > max_mb * 1024 * 1024:
            return func.HttpResponse(f"File too large (>{max_mb}MB)", status_code=413)
        safe_name = _sanitize_filename(req.params.get("fileName") or filename)
        resolved_ct = _guess_mime_type(safe_name, content_type)
        # Normalize extension; strip quotes and spaces
        _, ext = os.path.splitext(safe_name.strip().strip('"').lower())
        if resolved_ct not in _ALLOWED_TEMPLATE_MIME_TYPES and ext not in _ALLOWED_TEMPLATE_EXTENSIONS:
            return func.HttpResponse("Unsupported template type (allowed: .docx, .dotx)", status_code=415)
        blob_name = f"{user_id}/templates/{safe_name}"
        blob_client = _get_blob_client(blob_name)
        cs = _make_content_settings(resolved_ct)
        _upload_bytes_to_blob(blob_client, data, content_settings=cs)
        props = blob_client.get_blob_properties()
        meta = {
            "blob": blob_name,
            "etag": getattr(props, "etag", None),
            "size": getattr(props, "size", None),
            "contentType": getattr(props, "content_settings", None).content_type if getattr(props, "content_settings", None) else resolved_ct,
            "lastModified": getattr(props, "last_modified", None).isoformat() if getattr(props, "last_modified", None) else None,
        }
        sas = _generate_blob_sas_url(blob_name, permissions="r")
        meta["sasUrl"] = sas.get("url")
        meta["expiresUtc"] = sas.get("expiresUtc")
        return func.HttpResponse(body=json.dumps(meta), mimetype="application/json", status_code=201)
    except Exception as exc:
        logging.exception("upload_template failed")
        return func.HttpResponse(f"upload failed: {exc}", status_code=500)


@app.route(route="users/templates", methods=["GET"], auth_level=func.AuthLevel.FUNCTION)
def list_templates_http(req: func.HttpRequest) -> func.HttpResponse:
    user_id = None
    try:
        data = req.get_json()
        if isinstance(data, dict):
            user_id = data.get("user_id") or data.get("userId")
    except Exception:
        pass
    if not user_id:
        return func.HttpResponse("Missing user_id in JSON body", status_code=400)
    _init_storage()
    container_client = _blob_service_client.get_container_client(
        _blob_container_name)
    prefix = f"{user_id}/templates/"
    try:
        page_size_val: int = 50
        if isinstance(data, dict) and data.get("pageSize") is not None:
            try:
                page_size_val = int(data.get("pageSize"))
            except Exception:
                page_size_val = 50
        page_size = max(1, min(page_size_val, 200))
        token = req.params.get("continuationToken")
        include_shared = (req.params.get(
            "includeShared", "false").lower() == "true")
        # Page user templates
        pager_user = container_client.list_blobs(
            name_starts_with=prefix, results_per_page=page_size).by_page(continuation_token=token)
        user_items = []
        user_token = None
        for page in pager_user:
            for blob in page:
                if blob.name.endswith("/.keep"):
                    continue
                user_items.append({
                    "name": blob.name,
                    "size": getattr(blob, "size", None),
                    "contentType": getattr(getattr(blob, "content_settings", None), "content_type", None),
                    "lastModified": getattr(blob, "last_modified", None).isoformat() if getattr(blob, "last_modified", None) else None,
                })
            user_token = getattr(pager_user, "continuation_token", None)
            break
        payload = {"items": user_items}
        if user_token:
            payload["continuationToken"] = user_token
        if include_shared:
            # Page shared templates separately
            pager_shared = container_client.list_blobs(name_starts_with=_ORG_TEMPLATES_PREFIX, results_per_page=page_size).by_page(
                continuation_token=req.params.get("sharedContinuationToken"))
            shared_items = []
            shared_token = None
            for page in pager_shared:
                for blob in page:
                    if blob.name.endswith("/.keep"):
                        continue
                    shared_items.append({
                        "name": blob.name,
                        "size": getattr(blob, "size", None),
                        "contentType": getattr(getattr(blob, "content_settings", None), "content_type", None),
                        "lastModified": getattr(blob, "last_modified", None).isoformat() if getattr(blob, "last_modified", None) else None,
                    })
                shared_token = getattr(
                    pager_shared, "continuation_token", None)
                break
            payload["shared"] = {"items": shared_items}
            if shared_token:
                payload["shared"]["continuationToken"] = shared_token
        return func.HttpResponse(body=json.dumps(payload), mimetype="application/json", status_code=200)
    except Exception as exc:
        logging.exception("list_templates_http failed")
        return func.HttpResponse(f"list failed: {exc}", status_code=500)


@app.route(route="templates", methods=["GET"], auth_level=func.AuthLevel.FUNCTION)
def list_shared_templates(req: func.HttpRequest) -> func.HttpResponse:
    _init_storage()
    container_client = _blob_service_client.get_container_client(
        _blob_container_name)
    prefix = _ORG_TEMPLATES_PREFIX
    try:
        try:
            page_size = int(req.params.get("pageSize", "50"))
        except Exception:
            page_size = 50
        page_size = max(1, min(page_size, 200))
        token = req.params.get("continuationToken")
        pager = container_client.list_blobs(
            name_starts_with=prefix, results_per_page=page_size).by_page(continuation_token=token)
        items = []
        next_token = None
        for page in pager:
            for blob in page:
                if blob.name.endswith("/.keep"):
                    continue
                items.append({
                    "name": blob.name,
                    "size": getattr(blob, "size", None),
                    "contentType": getattr(getattr(blob, "content_settings", None), "content_type", None),
                    "lastModified": getattr(blob, "last_modified", None).isoformat() if getattr(blob, "last_modified", None) else None,
                })
            next_token = getattr(pager, "continuation_token", None)
            break
        payload = {"items": items}
        if next_token:
            payload["continuationToken"] = next_token
        return func.HttpResponse(body=json.dumps(payload), mimetype="application/json", status_code=200)
    except Exception as exc:
        logging.exception("list_shared_templates failed")
        return func.HttpResponse(f"list failed: {exc}", status_code=500)


# Removed: promote_template endpoint (renaming/copy deprecated)


word_tool_props_add_paragraph = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx blob name in the container (without user prefix)").to_dict(),
        ToolProperty("text", "string", "Paragraph text to append").to_dict(),
        ToolProperty("style", "string", "Optional Word style name").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_add_paragraph",
    description="Append a paragraph to an existing Word document stored in /tmp.",
    toolProperties=word_tool_props_add_paragraph,
)
def word_add_paragraph(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    text = args.get("text", "")
    style = args.get("style")
    if not filename:
        return "Missing filename"
    qualified_blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(qualified_blob_name)
    doc = _docx.Document(local_path)
    if style:
        doc.add_paragraph(text, style=style)
    else:
        doc.add_paragraph(text)
    doc.save(local_path)
    _upload_file_to_blob(local_path, qualified_blob_name)
    sas = _generate_blob_sas_url(qualified_blob_name, permissions="r")
    return json.dumps({"blob": qualified_blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_get_text = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_get_text",
    description="Get all text from a Word document stored in /tmp.",
    toolProperties=word_tool_props_get_text,
)
def word_get_text(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    if not filename:
        return "Missing filename"
    qualified_blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(qualified_blob_name)
    doc = _docx.Document(local_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text


# Additional Word tools
word_tool_props_add_heading = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("text", "string", "Heading text").to_dict(),
        ToolProperty("level", "number", "Heading level 1-9").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_add_heading",
    description="Add a heading to a Word document stored in Blob.",
    toolProperties=word_tool_props_add_heading,
)
def word_add_heading(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    text = args.get("text", "")
    level = int(args.get("level", 1))
    if not filename:
        return "Missing filename"
    qualified_blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(qualified_blob_name)
    doc = _docx.Document(local_path)
    doc.add_heading(text, level=level)
    doc.save(local_path)
    _upload_file_to_blob(local_path, qualified_blob_name)
    sas = _generate_blob_sas_url(qualified_blob_name, permissions="r")
    return json.dumps({"blob": qualified_blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_copy = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("source_filename", "string",
                     "Source .docx filename (without user prefix)").to_dict(),
        ToolProperty("destination_filename", "string",
                     "Destination .docx filename (without user prefix; optional; UUID if omitted)").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_copy_document",
    description="Copy a Word document blob to a new blob name.",
    toolProperties=word_tool_props_copy,
)
def word_copy_document(context) -> str:
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    source = args.get("source_filename")
    destination = args.get("destination_filename") or f"{uuid.uuid4()}.docx"
    if not source:
        return "Missing source_filename"
    src_blob = _get_blob_client(f"{user_id}/{source}" if user_id else source)
    dest_blob = _get_blob_client(
        f"{user_id}/{destination}" if user_id else destination)
    # Server-side copy
    dest_blob.start_copy_from_url(src_blob.url)
    return f"Copied to blob '{_blob_container_name}/" + (f"{user_id}/{destination}" if user_id else destination) + "'"


word_tool_props_search_replace = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("find_text", "string", "Text to find").to_dict(),
        ToolProperty("replace_text", "string",
                     "Text to replace with").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_search_and_replace",
    description="Search and replace text in a Word document (paragraph-level, naive).",
    toolProperties=word_tool_props_search_replace,
)
def word_search_and_replace(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    find_text = args.get("find_text", "")
    replace_text = args.get("replace_text", "")
    if not filename:
        return "Missing filename"
    if not find_text:
        return "Missing find_text"
    qualified_blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(qualified_blob_name)
    doc = _docx.Document(local_path)
    for paragraph in doc.paragraphs:
        if find_text in paragraph.text:
            paragraph.text = paragraph.text.replace(find_text, replace_text)
    doc.save(local_path)
    _upload_file_to_blob(local_path, qualified_blob_name)
    sas = _generate_blob_sas_url(qualified_blob_name, permissions="r")
    return json.dumps({"blob": qualified_blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


# List documents in the container
word_tool_props_list_docs = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_list_documents",
    description="List .docx documents available in the blob container.",
    toolProperties=word_tool_props_list_docs,
)
def word_list_documents(context) -> str:
    _init_storage()
    container_client = _blob_service_client.get_container_client(
        _blob_container_name)
    names = []
    try:
        payload = json.loads(context)
        args = payload.get("arguments", {})
        user_id = args.get("user_id")
        prefix = f"{user_id}/" if user_id else ""
        for blob in container_client.list_blobs(name_starts_with=prefix):
            if blob.name.lower().endswith(".docx"):
                names.append(blob.name)
    except Exception as exc:
        return f"Failed to list blobs: {exc}"
    return json.dumps(names)


# Document info
word_tool_props_get_info = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_get_document_info",
    description="Get core document properties (title, author).",
    toolProperties=word_tool_props_get_info,
)
def word_get_document_info(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    if not filename:
        return "Missing filename"
    qualified_blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(qualified_blob_name)
    doc = _docx.Document(local_path)
    props = doc.core_properties
    info = {
        "title": getattr(props, "title", None),
        "author": getattr(props, "author", None),
        "created": str(getattr(props, "created", None)),
        "last_modified_by": getattr(props, "last_modified_by", None),
    }
    return json.dumps(info)


# Document outline (headings)
word_tool_props_get_outline = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_get_document_outline",
    description="Get document outline by collecting heading paragraphs.",
    toolProperties=word_tool_props_get_outline,
)
def word_get_document_outline(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    if not filename:
        return "Missing filename"
    qualified_blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(qualified_blob_name)
    doc = _docx.Document(local_path)
    outline = []
    for p in doc.paragraphs:
        style_name = getattr(getattr(p, "style", None), "name", "") or ""
        if style_name.lower().startswith("heading"):
            outline.append({"text": p.text, "style": style_name})
    return json.dumps(outline)


# HTTP dispatcher for Copilot Studio integration
@app.route(route="mcp/exec", methods=["POST"], auth_level=func.AuthLevel.FUNCTION)
def mcp_exec(req: func.HttpRequest) -> func.HttpResponse:
    try:
        body = req.get_json()
    except Exception:
        return func.HttpResponse("Invalid JSON body", status_code=400)
    tool_name = body.get("toolName")
    arguments = body.get("arguments", {})
    if not tool_name:
        return func.HttpResponse("Missing toolName", status_code=400)
    # Map tool names to callables defined above
    tools = {
        "hello_mcp": hello_mcp,
        "word_create_document": word_create_document,
        "word_add_paragraph": word_add_paragraph,
        "word_get_text": word_get_text,
        "word_add_heading": word_add_heading,
        "word_copy_document": word_copy_document,
        "word_search_and_replace": word_search_and_replace,
        "word_list_documents": word_list_documents,
        "word_list_templates": word_list_templates,
        "word_get_document_info": word_get_document_info,
        "word_get_document_outline": word_get_document_outline,
        "word_add_table": word_add_table,
        "word_add_picture": word_add_picture,
        "word_add_page_break": word_add_page_break,
        "word_get_paragraph_text": word_get_paragraph_text,
        "word_find_text": word_find_text,
        "word_format_text": word_format_text,
        "word_delete_paragraph": word_delete_paragraph,
        "word_create_custom_style": word_create_custom_style,
        "word_format_table": word_format_table,
        "word_set_table_cell_shading": word_set_table_cell_shading,
        "word_apply_table_alternating_rows": word_apply_table_alternating_rows,
        "word_highlight_table_header": word_highlight_table_header,
        "word_merge_table_cells": word_merge_table_cells,
        "word_merge_table_cells_horizontal": word_merge_table_cells_horizontal,
        "word_merge_table_cells_vertical": word_merge_table_cells_vertical,
        "word_set_table_cell_alignment": word_set_table_cell_alignment,
        "word_set_table_alignment_all": word_set_table_alignment_all,
        "word_format_table_cell_text": word_format_table_cell_text,
        "word_set_table_cell_padding": word_set_table_cell_padding,
        "word_set_table_column_width": word_set_table_column_width,
        "word_set_table_width": word_set_table_width,
        "word_auto_fit_table_columns": word_auto_fit_table_columns,
        "word_get_all_comments": word_get_all_comments,
        "word_get_comments_by_author": word_get_comments_by_author,
        "word_get_comments_for_paragraph": word_get_comments_for_paragraph,
        "word_add_comment": word_add_comment,
    }
    func_to_call = tools.get(tool_name)
    if not func_to_call:
        return func.HttpResponse(f"Unknown tool: {tool_name}", status_code=404)
    try:
        result = func_to_call(json.dumps({"arguments": arguments}))
    except Exception as exc:
        logging.exception("Tool execution failed")
        return func.HttpResponse(f"Tool execution error: {exc}", status_code=500)
    if isinstance(result, (dict, list)):
        body = json.dumps(result)
    else:
        body = str(result)
    return func.HttpResponse(body=body, mimetype="application/json", status_code=200)


@app.route(route="convert/word-to-pdf", methods=["POST"], auth_level=func.AuthLevel.FUNCTION)
def convert_word_to_pdf(req: func.HttpRequest) -> func.HttpResponse:
    """Temporary gateway: upload a Word file to SharePoint via Graph, convert to PDF, return PDF, cleanup."""
    try:
        # Accept either: (a) direct upload multipart/octet-stream; (b) blob path
        blob_path = req.params.get("blob")
        filename = req.params.get("fileName")
        data = None
        if blob_path:
            _init_storage()
            local = _download_blob_to_temp(blob_path)
            with open(local, "rb") as f:
                data = f.read()
            if not filename:
                filename = os.path.basename(blob_path)
        else:
            data, filename_in, ct = _read_form_file(req, field_name="file")
            if not data:
                return func.HttpResponse("Empty body", status_code=400)
            filename = filename or filename_in or "document.docx"
        # Ensure extension docx
        base, ext = os.path.splitext(filename)
        if ext.lower() not in (".docx", ".dotx"):
            return func.HttpResponse("Only .docx or .dotx supported", status_code=415)
        drive_id = _resolve_sharepoint_drive()
        # Use original name without UUID (temporary file is deleted after conversion)
        remote_name = f"tmp/{base}{ext}"
        # Upload with same extension (.docx or .dotx)
        item = _graph_upload_to_drive(drive_id, remote_name, data)
        item_id = item.get("id")
        if not item_id:
            raise RuntimeError("Upload did not return item id")
        try:
            pdf_bytes = _graph_download_pdf_content(drive_id, item_id)
        finally:
            try:
                _graph_delete_item(drive_id, item_id)
            except Exception as exc:
                logging.warning(
                    "Failed to delete temporary item %s in drive %s: %s",
                    item_id,
                    drive_id,
                    exc,
                )
        # Store PDF to Blob and return SAS URL
        _init_storage()
        # Compute destination blob path
        try:
            body_json = req.get_json()
        except Exception:
            body_json = {}
        explicit_dest = (body_json.get("dest") if isinstance(
            body_json, dict) else None) or req.params.get("dest")
        user_id = req.params.get("user_id") or (body_json.get(
            "user_id") if isinstance(body_json, dict) else None)
        base_name = os.path.splitext(os.path.basename(filename))[0]
        pdf_file = f"{base_name}.pdf"
        if explicit_dest:
            if "/" in explicit_dest:
                dest_blob_name = explicit_dest
            else:
                # If blob_path present, place under same directory; else under user scope
                if blob_path:
                    dest_blob_name = f"{os.path.dirname(blob_path).rstrip('/')}/{_sanitize_filename(explicit_dest)}"
                elif user_id:
                    dest_blob_name = f"{user_id}/{_sanitize_filename(explicit_dest)}"
                else:
                    return func.HttpResponse("Missing user_id to resolve dest", status_code=400)
        else:
            if blob_path:
                parent = os.path.dirname(blob_path).rstrip("/")
                dest_blob_name = f"{parent}/{_sanitize_filename(pdf_file)}" if parent else _sanitize_filename(
                    pdf_file)
            else:
                if not user_id:
                    return func.HttpResponse("Missing user_id for upload destination", status_code=400)
                dest_blob_name = f"{user_id}/{_sanitize_filename(pdf_file)}"
        blob_client = _get_blob_client(dest_blob_name)
        cs = _make_content_settings("application/pdf")

        def _op():
            return blob_client.upload_blob(pdf_bytes, overwrite=True, content_settings=cs)
        _with_retries(_op)
        _apply_blob_ttl(blob_client)
        props = blob_client.get_blob_properties()
        meta = {
            "container": _blob_container_name,
            "blob": dest_blob_name,
            "etag": getattr(props, "etag", None),
            "size": getattr(props, "size", None),
            "contentType": getattr(props, "content_settings", None).content_type if getattr(props, "content_settings", None) else "application/pdf",
            "lastModified": getattr(props, "last_modified", None).isoformat() if getattr(props, "last_modified", None) else None,
        }
        sas = _generate_blob_sas_url(dest_blob_name, permissions="r")
        meta["sasUrl"] = sas.get("url")
        meta["expiresUtc"] = sas.get("expiresUtc")
        return func.HttpResponse(body=json.dumps(meta), mimetype="application/json", status_code=200)
    except Exception as exc:
        logging.exception("convert_word_to_pdf failed")
        return func.HttpResponse(f"conversion failed: {exc}", status_code=500)


# ---- Vague 2: Advanced table formatting tools ----

def _ensure_cell_shading(cell, fill_color: str, pattern: str = "clear"):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), pattern)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(fill_color))
    tc_pr.append(shd)


def _set_cell_padding(cell, top=None, bottom=None, left=None, right=None, unit=None):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    def to_twips(val):
        try:
            return str(int(float(val) * 20))
        except Exception:
            return None
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in ("top", top), ("bottom", bottom), ("left", left), ("right", right):
        if value is None:
            continue
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        tw = to_twips(value)
        if tw is not None:
            el.set(qn("w:w"), tw)
            el.set(qn("w:type"), "dxa")


def _set_table_width(table, width_value: float, width_type: str = "points"):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    def to_twips(val):
        try:
            return str(int(float(val) * 20))
        except Exception:
            return None
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    if (width_type or "").lower() == "percent":
        try:
            pct = str(int(float(width_value) * 50))
            tbl_w.set(qn("w:w"), pct)
            tbl_w.set(qn("w:type"), "pct")
        except Exception as exc:
            logging.warning(
                "Failed to set table width to %s (file: unknown): %s",
                width_value,
                exc,
            )
    else:
        tw = to_twips(width_value)
        if tw is not None:
            tbl_w.set(qn("w:w"), tw)
            tbl_w.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_value: float, width_type: str = "points"):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    def to_twips(val):
        try:
            return str(int(float(val) * 20))
        except Exception:
            return None
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    if (width_type or "").lower() == "percent":
        try:
            pct = str(int(float(width_value) * 50))
            tc_w.set(qn("w:w"), pct)
            tc_w.set(qn("w:type"), "pct")
        except Exception as exc:
            logging.warning(
                "Failed to set cell width to %s (file: unknown): %s",
                width_value,
                exc,
            )
    else:
        tw = to_twips(width_value)
        if tw is not None:
            tc_w.set(qn("w:w"), tw)
            tc_w.set(qn("w:type"), "dxa")


def _set_table_layout_autofit(table):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")
    try:
        table.autofit = True
    except Exception as exc:
        logging.warning(
            "Failed to enable table autofit (file: unknown): %s",
            exc,
        )


word_tool_props_format_table = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("has_header_row", "boolean",
                 "Mark first row as header (visual only)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_format_table",
    description="Basic table formatting (header flag only).",
    toolProperties=word_tool_props_format_table,
)
def word_format_table(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    table_index = int(args.get("table_index", -1))
    has_header_row = bool(args.get("has_header_row", False))
    if not filename or table_index < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if table_index >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[table_index]
    if has_header_row and len(table.rows) > 0:
        try:
            for p in table.rows[0].cells[0].paragraphs:
                logging.warning(
                    "No operation for header paragraph in file %s (user %s)",
                    blob_name,
                    user_id,
                )
            # Visual cue: bold header row
            for cell in table.rows[0].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        except Exception as exc:
            logging.warning(
                "Failed to format header row for %s (user %s): %s",
                blob_name,
                user_id,
                exc,
            )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_cell_shading = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("fill_color", "string", "Hex color without #").to_dict(),
    ToolProperty("pattern", "string", "Pattern (default clear)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_cell_shading",
    description="Set shading on a given table cell.",
    toolProperties=word_tool_props_set_table_cell_shading,
)
def word_set_table_cell_shading(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r_idx = int(args.get("row_index", -1))
    c_idx = int(args.get("col_index", -1))
    fill_color = args.get("fill_color")
    pattern = args.get("pattern", "clear")
    if not filename or t_idx < 0 or r_idx < 0 or c_idx < 0 or not fill_color:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if r_idx >= len(table.rows) or c_idx >= len(table.columns):
        return "cell index out of range"
    _ensure_cell_shading(table.cell(r_idx, c_idx), fill_color, pattern)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_apply_table_alternating_rows = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("color1", "string", "Hex color 1 (default FFFFFF)").to_dict(),
    ToolProperty("color2", "string", "Hex color 2 (default F2F2F2)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_apply_table_alternating_rows",
    description="Apply alternating row colors.",
    toolProperties=word_tool_props_apply_table_alternating_rows,
)
def word_apply_table_alternating_rows(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    color1 = args.get("color1", "FFFFFF")
    color2 = args.get("color2", "F2F2F2")
    if not filename or t_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    for r_idx, row in enumerate(table.rows):
        fill = color1 if r_idx % 2 == 0 else color2
        for c_idx, _ in enumerate(table.columns):
            _ensure_cell_shading(table.cell(r_idx, c_idx), fill, "clear")
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_highlight_table_header = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("header_color", "string",
                 "Header color (default 4472C4)").to_dict(),
    ToolProperty("text_color", "string",
                 "Header text color (default FFFFFF)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_highlight_table_header",
    description="Highlight first row as header with colors.",
    toolProperties=word_tool_props_highlight_table_header,
)
def word_highlight_table_header(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    header_color = args.get("header_color", "4472C4")
    text_color = args.get("text_color", "FFFFFF")
    if not filename or t_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables) or len(doc.tables[t_idx].rows) == 0:
        return "table_index out of range"
    table = doc.tables[t_idx]
    for cell in table.rows[0].cells:
        _ensure_cell_shading(cell, header_color, "clear")
        try:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = _docx.shared.RGBColor.from_string(text_color)
        except Exception:
            pass
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_merge_cells = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("start_row", "number", "Start row index").to_dict(),
    ToolProperty("start_col", "number", "Start column index").to_dict(),
    ToolProperty("end_row", "number", "End row index").to_dict(),
    ToolProperty("end_col", "number", "End column index").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_merge_table_cells",
    description="Merge a rectangular range of cells.",
    toolProperties=word_tool_props_merge_cells,
)
def word_merge_table_cells(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    sr = int(args.get("start_row", -1))
    sc = int(args.get("start_col", -1))
    er = int(args.get("end_row", -1))
    ec = int(args.get("end_col", -1))
    if not filename or t_idx < 0 or sr < 0 or sc < 0 or er < 0 or ec < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    try:
        table.cell(sr, sc).merge(table.cell(er, ec))
    except Exception as exc:
        return f"merge failed: {exc}"
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_merge_cells_horizontal = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index").to_dict(),
    ToolProperty("start_col", "number", "Start column index").to_dict(),
    ToolProperty("end_col", "number", "End column index").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_merge_table_cells_horizontal",
    description="Merge horizontally in a single row.",
    toolProperties=word_tool_props_merge_cells_horizontal,
)
def word_merge_table_cells_horizontal(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r = int(args.get("row_index", -1))
    sc = int(args.get("start_col", -1))
    ec = int(args.get("end_col", -1))
    if not filename or t_idx < 0 or r < 0 or sc < 0 or ec < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    try:
        table.cell(r, sc).merge(table.cell(r, ec))
    except Exception as exc:
        return f"merge failed: {exc}"
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_merge_cells_vertical = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index").to_dict(),
    ToolProperty("start_row", "number", "Start row index").to_dict(),
    ToolProperty("end_row", "number", "End row index").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_merge_table_cells_vertical",
    description="Merge vertically in a single column.",
    toolProperties=word_tool_props_merge_cells_vertical,
)
def word_merge_table_cells_vertical(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    c = int(args.get("col_index", -1))
    sr = int(args.get("start_row", -1))
    er = int(args.get("end_row", -1))
    if not filename or t_idx < 0 or c < 0 or sr < 0 or er < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    try:
        table.cell(sr, c).merge(table.cell(er, c))
    except Exception as exc:
        return f"merge failed: {exc}"
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_cell_alignment = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("horizontal", "string",
                 "left|center|right|justify").to_dict(),
    ToolProperty("vertical", "string", "top|center|bottom").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_cell_alignment",
    description="Set horizontal/vertical alignment for a cell.",
    toolProperties=word_tool_props_set_table_cell_alignment,
)
def word_set_table_cell_alignment(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r = int(args.get("row_index", -1))
    c = int(args.get("col_index", -1))
    horizontal = (args.get("horizontal") or "").lower()
    vertical = (args.get("vertical") or "").lower()
    if not filename or t_idx < 0 or r < 0 or c < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if r >= len(table.rows) or c >= len(table.columns):
        return "cell index out of range"
    cell = table.cell(r, c)
    try:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if horizontal:
            mapping = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            align = mapping.get(horizontal)
            if align is not None:
                for p in cell.paragraphs:
                    p.alignment = align
        if vertical:
            vm = {
                "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
                "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
            }
            v = vm.get(vertical)
            if v is not None:
                cell.vertical_alignment = v
    except Exception as exc:
        logging.warning(
            "Failed to set cell alignment for %s (user %s): %s",
            blob_name,
            user_id,
            exc,
        )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_alignment_all = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("horizontal", "string",
                 "left|center|right|justify").to_dict(),
    ToolProperty("vertical", "string", "top|center|bottom").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_alignment_all",
    description="Apply alignment to all cells in a table.",
    toolProperties=word_tool_props_set_table_alignment_all,
)
def word_set_table_alignment_all(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    horizontal = (args.get("horizontal") or "").lower()
    vertical = (args.get("vertical") or "").lower()
    if not filename or t_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    try:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        hmap = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        vmap = {
            "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
            "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        }
        halign = hmap.get(horizontal)
        valign = vmap.get(vertical)
        for row in table.rows:
            for cell in row.cells:
                if halign is not None:
                    for p in cell.paragraphs:
                        p.alignment = halign
                if valign is not None:
                    cell.vertical_alignment = valign
    except Exception as exc:
        logging.warning(
            "Failed to set table alignment for %s (user %s): %s",
            blob_name,
            user_id,
            exc,
        )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_format_table_cell_text = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("text_content", "string", "Replace cell text").to_dict(),
    ToolProperty("bold", "boolean", "Bold").to_dict(),
    ToolProperty("italic", "boolean", "Italic").to_dict(),
    ToolProperty("underline", "boolean", "Underline").to_dict(),
    ToolProperty("color", "string", "Hex color without #").to_dict(),
    ToolProperty("font_size", "number", "Font size in points").to_dict(),
    ToolProperty("font_name", "string", "Font family name").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_format_table_cell_text",
    description="Format text in a specific table cell.",
    toolProperties=word_tool_props_format_table_cell_text,
)
def word_format_table_cell_text(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r = int(args.get("row_index", -1))
    c = int(args.get("col_index", -1))
    text = args.get("text_content")
    bold = args.get("bold")
    italic = args.get("italic")
    underline = args.get("underline")
    color = args.get("color")
    font_size = args.get("font_size")
    font_name = args.get("font_name")
    if not filename or t_idx < 0 or r < 0 or c < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if r >= len(table.rows) or c >= len(table.columns):
        return "cell index out of range"
    cell = table.cell(r, c)
    if text is not None:
        cell.text = str(text)
    try:
        for p in cell.paragraphs:
            for run in p.runs:
                if bold is not None:
                    run.bold = bool(bold)
                if italic is not None:
                    run.italic = bool(italic)
                if underline is not None:
                    run.underline = bool(underline)
                if color:
                    try:
                        run.font.color.rgb = _docx.shared.RGBColor.from_string(str(color))
                    except Exception:
                        pass
                if font_size is not None:
                    try:
                        run.font.size = _docx.shared.Pt(float(font_size))
                    except Exception as exc:
                        logging.warning(
                            "Failed to set font size for %s (user %s): %s",
                            blob_name,
                            user_id,
                            exc,
                        )
                if font_name:
                    run.font.name = str(font_name)
    except Exception as exc:
        logging.error(
            "Failed to format table cell text for %s (user %s): %s",
            blob_name,
            user_id,
            exc,
        )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_cell_padding = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("row_index", "number", "Row index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("top", "number", "Top padding (points)").to_dict(),
    ToolProperty("bottom", "number", "Bottom padding (points)").to_dict(),
    ToolProperty("left", "number", "Left padding (points)").to_dict(),
    ToolProperty("right", "number", "Right padding (points)").to_dict(),
    ToolProperty("unit", "string",
                 "Padding unit in points (only 'points' supported)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_cell_padding",
    description="Set padding for a specific table cell.",
    toolProperties=word_tool_props_set_table_cell_padding,
)
def word_set_table_cell_padding(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    r = int(args.get("row_index", -1))
    c = int(args.get("col_index", -1))
    top = args.get("top")
    bottom = args.get("bottom")
    left = args.get("left")
    right = args.get("right")
    unit = args.get("unit") or "points"
    if not filename or t_idx < 0 or r < 0 or c < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if r >= len(table.rows) or c >= len(table.columns):
        return "cell index out of range"
    _set_cell_padding(table.cell(r, c), top=top, bottom=bottom,
                      left=left, right=right, unit=unit)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_column_width = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("col_index", "number", "Column index (0-based)").to_dict(),
    ToolProperty("width", "number",
                 "Width value (points or percent)").to_dict(),
    ToolProperty("width_type", "string",
                 "points|percent (default points)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_column_width",
    description="Set width for a specific column.",
    toolProperties=word_tool_props_set_table_column_width,
)
def word_set_table_column_width(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    c = int(args.get("col_index", -1))
    width = args.get("width")
    width_type = args.get("width_type") or "points"
    if not filename or t_idx < 0 or c < 0 or width is None:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    if c >= len(table.columns):
        return "col_index out of range"
    for row in table.rows:
        _set_cell_width(row.cells[c], width, width_type)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_set_table_column_widths = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("widths", "array", "Array of widths by column (points unless width_type=percent)",
                 item_type="number").to_dict(),
    ToolProperty("width_type", "string",
                 "points|percent (default points)").to_dict(),
])


word_tool_props_set_table_width = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
    ToolProperty("width", "number", "Table width value").to_dict(),
    ToolProperty("width_type", "string",
                 "points|percent (default points)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_set_table_width",
    description="Set overall table width.",
    toolProperties=word_tool_props_set_table_width,
)
def word_set_table_width(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    width = args.get("width")
    width_type = args.get("width_type") or "points"
    if not filename or t_idx < 0 or width is None:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    _set_table_width(table, width, width_type)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_auto_fit_table_columns = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string", "Target .docx filename").to_dict(),
    ToolProperty("table_index", "number", "Table index (0-based)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_auto_fit_table_columns",
    description="Enable table auto-fit layout.",
    toolProperties=word_tool_props_auto_fit_table_columns,
)
def word_auto_fit_table_columns(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    t_idx = int(args.get("table_index", -1))
    if not filename or t_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if t_idx >= len(doc.tables):
        return "table_index out of range"
    table = doc.tables[t_idx]
    _set_table_layout_autofit(table)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


# ---- Vague 3: Comments and PDF conversion ----

word_tool_props_add_comment = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string",
                 "Target .docx filename (without user prefix)").to_dict(),
    ToolProperty("paragraph_index", "number",
                 "Zero-based paragraph index to attach comment to").to_dict(),
    ToolProperty("text", "string", "Comment text").to_dict(),
    ToolProperty("author", "string", "Optional author name").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_add_comment",
    description="Add a comment to a paragraph (best-effort via XML).",
    toolProperties=word_tool_props_add_comment,
)
def word_add_comment(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    p_idx = int(args.get("paragraph_index", -1))
    text = args.get("text")
    author = args.get("author") or "Author"
    if not filename or p_idx < 0 or not text:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if p_idx >= len(doc.paragraphs):
        return "paragraph_index out of range"
    p = doc.paragraphs[p_idx]
    # Attach comment via comments part (WordprocessingML). python-docx n'a pas d'API publique.
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        part = doc.part
        # Ensure comments part exists
        rels = [r for r in part.rels.values() if r.reltype == RT.COMMENTS]
        if rels:
            comments_part = rels[0].target_part
        else:
            return "Comments part not found; create the document from a template that already contains comments (even empty) and retry"
        # Parse and append a new <w:comment>
        from xml.etree import ElementTree as ET
        xml = comments_part.blob.decode("utf-8", errors="ignore")
        root = ET.fromstring(xml)
        # Compute new ID
        existing_ids = [int(el.attrib.get(qn("w:id"), "0")) for el in root.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment")]
        new_id = max(existing_ids) + 1 if existing_ids else 0
        c_el = ET.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment", {
            qn("w:id"): str(new_id),
            qn("w:author"): author,
        })
        p_el = ET.SubElement(
            c_el, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
        r_el = ET.SubElement(
            p_el, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        t_el = ET.SubElement(
            r_el, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t_el.text = str(text)
        root.append(c_el)
        comments_part._blob = ET.tostring(
            root, encoding="utf-8", xml_declaration=True)
        # Mark paragraph with comment range start/end
        p_elm = p._element
        start = OxmlElement("w:commentRangeStart")
        start.set(qn("w:id"), str(new_id))
        end = OxmlElement("w:commentRangeEnd")
        end.set(qn("w:id"), str(new_id))
        ref = OxmlElement("w:r")
        cref = OxmlElement("w:commentReference")
        cref.set(qn("w:id"), str(new_id))
        ref.append(cref)
        p_elm.insert(0, start)
        p_elm.append(end)
        p_elm.append(ref)
    except Exception as exc:
        return f"Failed to add comment: {exc}"
    # Save back
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc"), "commentId": new_id})


word_tool_props_get_all_comments = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string",
                 "Target .docx filename (without user prefix)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_get_all_comments",
    description="Extract all comments from the document.",
    toolProperties=word_tool_props_get_all_comments,
)
def word_get_all_comments(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    if not filename:
        return "Missing filename"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    comments = []
    try:
        # python-docx n'a pas d'API publique pour commentaires; on lit l'XML
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        part = doc.part
        rels = [r for r in part.rels.values() if r.reltype == RT.COMMENTS]
        if rels:
            comments_part = rels[0].target_part
            xml = comments_part.blob.decode("utf-8", errors="ignore")
            # Minimal parsing: extraire <w:t> dans w:comment
            import re
            for cid, ctext in re.findall(r"<w:comment[\s\S]*?w:id=\"(\d+)\"[\s\S]*?<w:t>([\s\S]*?)</w:t>", xml):
                comments.append({"id": cid, "text": ctext})
    except Exception as exc:
        logging.warning(
            "Failed to extract comments for %s (user %s): %s",
            blob_name,
            user_id,
            exc,
        )
    return json.dumps(comments)


word_tool_props_get_comments_by_author = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string",
                 "Target .docx filename (without user prefix)").to_dict(),
    ToolProperty("author", "string", "Author name to filter").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_get_comments_by_author",
    description="Extract comments by author.",
    toolProperties=word_tool_props_get_comments_by_author,
)
def word_get_comments_by_author(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    author = args.get("author")
    if not filename or not author:
        return "Missing filename or author"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    results = []
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        part = doc.part
        rels = [r for r in part.rels.values() if r.reltype == RT.COMMENTS]
        if rels:
            comments_part = rels[0].target_part
            xml = comments_part.blob.decode("utf-8", errors="ignore")
            import re
            pattern = re.compile(
                rf"<w:comment[\s\S]*?w:author=\"{re.escape(author)}\"[\s\S]*?<w:t>([\s\S]*?)</w:t>")
            for match in pattern.findall(xml):
                results.append({"text": match})
    except Exception as exc:
        logging.warning(
            "Failed to extract comments by author %s for %s (user %s): %s",
            author,
            blob_name,
            user_id,
            exc,
        )
    return json.dumps(results)


word_tool_props_get_comments_for_paragraph = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string",
                 "Target .docx filename (without user prefix)").to_dict(),
    ToolProperty("paragraph_index", "number",
                 "Zero-based paragraph index").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_get_comments_for_paragraph",
    description="Get comments attached to a paragraph (best-effort via XML).",
    toolProperties=word_tool_props_get_comments_for_paragraph,
)
def word_get_comments_for_paragraph(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    p_idx = int(args.get("paragraph_index", -1))
    if not filename or p_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if p_idx >= len(doc.paragraphs):
        return "paragraph_index out of range"
    # Best effort: scan comments and look for paragraph text snippets
    para_text = doc.paragraphs[p_idx].text or ""
    matches = []
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        part = doc.part
        rels = [r for r in part.rels.values() if r.reltype == RT.COMMENTS]
        if rels:
            comments_part = rels[0].target_part
            xml = comments_part.blob.decode("utf-8", errors="ignore")
            import re
            for cid, ctext in re.findall(r"<w:comment[\s\S]*?w:id=\"(\d+)\"[\s\S]*?<w:t>([\s\S]*?)</w:t>", xml):
                if ctext and (ctext in para_text or para_text in ctext):
                    matches.append({"id": cid, "text": ctext})
    except Exception as exc:
        logging.warning(
            "Failed to get comments for paragraph %s in %s (user %s): %s",
            p_idx,
            blob_name,
            user_id,
            exc,
        )
    return json.dumps(matches)


word_tool_props_convert_to_pdf = json.dumps([
    ToolProperty("user_id", "string",
                 "User identifier used to namespace blobs").to_dict(),
    ToolProperty("filename", "string",
                 "Source .docx filename (without user prefix)").to_dict(),
    ToolProperty("pdf_filename", "string",
                 "Destination .pdf filename (without user prefix; optional)").to_dict(),
])

# ---- Storage utilities ----

word_tool_props_list_templates = json.dumps([
    ToolProperty("prefix", "string",
                 "Templates prefix (default templates/)").to_dict(),
])


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_list_templates",
    description="List available template blobs (default prefix templates/).",
    toolProperties=word_tool_props_list_templates,
)
def word_list_templates(context) -> str:
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    prefix = args.get("prefix") or "templates/"
    container_client = _blob_service_client.get_container_client(
        _blob_container_name)
    names = []
    try:
        for blob in container_client.list_blobs(name_starts_with=prefix):
            names.append(blob.name)
    except Exception as exc:
        return f"Failed to list templates: {exc}"
    return json.dumps(names)


# ---- Vague 1: Additional content and formatting tools ----

word_tool_props_add_table = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("rows", "number", "Number of rows").to_dict(),
        ToolProperty("cols", "number", "Number of columns").to_dict(),
        ToolProperty(
            "data",
            "string",
            "Optional JSON 2D array of cell texts, e.g. [[\"A\",\"B\"],[\"C\",\"D\"]]",
        ).to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_add_table",
    description="Insert a table; optionally populate with provided data.",
    toolProperties=word_tool_props_add_table,
)
def word_add_table(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    rows = int(args.get("rows", 0))
    cols = int(args.get("cols", 0))
    data = args.get("data")
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except Exception:
            data = None
    if not filename or rows <= 0 or cols <= 0:
        return "Missing filename or invalid rows/cols"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    table = doc.add_table(rows=rows, cols=cols)
    if isinstance(data, list):
        for r_idx in range(min(rows, len(data))):
            row_data = data[r_idx] if isinstance(data[r_idx], list) else []
            for c_idx in range(min(cols, len(row_data))):
                table.cell(r_idx, c_idx).text = str(row_data[c_idx])
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_add_picture = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("image_blob", "string",
                     "Path relative to container; supports subfolders (e.g., image_blob/watermark.png)").to_dict(),
        ToolProperty("width_points", "number",
                     "Optional width in points").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_add_picture",
    description="Insert a picture from a blob; optional width in points.",
    toolProperties=word_tool_props_add_picture,
)
def word_add_picture(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    image_blob = args.get("image_blob")
    width_points = args.get("width_points")
    if not filename or not image_blob:
        return "Missing filename or image_blob"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    # Resolve image blob path with tolerant search in common locations
    candidates = []
    # As provided
    candidates.append(image_blob)
    # Under user scope
    if user_id:
        candidates.append(f"{user_id}/{image_blob}")
    # Under default folder 'image_blob/'
    if not image_blob.startswith("image_blob/"):
        candidates.append(f"image_blob/{image_blob}")
        if user_id:
            candidates.append(f"{user_id}/image_blob/{image_blob}")
    image_blob_resolved = None
    for cand in candidates:
        if _blob_exists(cand):
            image_blob_resolved = cand
            break
    if not image_blob_resolved:
        return "Image blob not found"
    image_local = _download_blob_to_temp(image_blob_resolved)
    doc = _docx.Document(local_path)
    if width_points is not None:
        try:
            width = _docx.shared.Pt(float(width_points))
            doc.add_picture(image_local, width=width)
        except Exception:
            doc.add_picture(image_local)
    else:
        doc.add_picture(image_local)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_add_page_break = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_add_page_break",
    description="Insert a page break.",
    toolProperties=word_tool_props_add_page_break,
)
def word_add_page_break(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    if not filename:
        return "Missing filename"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    doc.add_page_break()
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_get_paragraph_text = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("paragraph_index", "number",
                     "Zero-based paragraph index").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_get_paragraph_text",
    description="Get text from a specific paragraph.",
    toolProperties=word_tool_props_get_paragraph_text,
)
def word_get_paragraph_text(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    paragraph_index = int(args.get("paragraph_index", -1))
    if not filename or paragraph_index < 0:
        return "Missing filename or invalid paragraph_index"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if paragraph_index >= len(doc.paragraphs):
        return "paragraph_index out of range"
    return doc.paragraphs[paragraph_index].text


word_tool_props_find_text = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("text_to_find", "string", "Text to search").to_dict(),
        ToolProperty("match_case", "boolean",
                     "Case sensitive (default true)").to_dict(),
        ToolProperty("whole_word", "boolean",
                     "Whole word match (default false)").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_find_text",
    description="Find text occurrences and return paragraph indexes.",
    toolProperties=word_tool_props_find_text,
)
def word_find_text(context) -> str:
    import re
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    text_to_find = args.get("text_to_find", "")
    match_case = bool(args.get("match_case", True))
    whole_word = bool(args.get("whole_word", False))
    if not filename or not text_to_find:
        return "Missing filename or text_to_find"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    flags = 0 if match_case else re.IGNORECASE
    pattern = re.escape(text_to_find)
    if whole_word:
        pattern = fr"\b{pattern}\b"
    regex = re.compile(pattern, flags)
    matches = []
    for idx, p in enumerate(doc.paragraphs):
        if regex.search(p.text or ""):
            matches.append({"paragraphIndex": idx})
    return json.dumps({"matches": matches})


word_tool_props_format_text = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("paragraph_index", "number",
                     "Zero-based paragraph index").to_dict(),
        ToolProperty("start_pos", "number",
                     "Start char index inclusive").to_dict(),
        ToolProperty("end_pos", "number",
                     "End char index exclusive").to_dict(),
        ToolProperty("bold", "boolean", "Set bold").to_dict(),
        ToolProperty("italic", "boolean", "Set italic").to_dict(),
        ToolProperty("underline", "boolean", "Set underline").to_dict(),
        ToolProperty("color", "string",
                     "Hex color without #, e.g., FF0000").to_dict(),
        ToolProperty("font_size", "number", "Font size in points").to_dict(),
        ToolProperty("font_name", "string", "Font family name").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_format_text",
    description="Format a text span within a paragraph (basic run-split).",
    toolProperties=word_tool_props_format_text,
)
def word_format_text(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    p_idx = int(args.get("paragraph_index", -1))
    start_pos = int(args.get("start_pos", -1))
    end_pos = int(args.get("end_pos", -1))
    bold = args.get("bold")
    italic = args.get("italic")
    underline = args.get("underline")
    color = args.get("color")
    font_size = args.get("font_size")
    font_name = args.get("font_name")
    if not filename or p_idx < 0 or start_pos < 0 or end_pos < 0 or end_pos <= start_pos:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if p_idx >= len(doc.paragraphs):
        return "paragraph_index out of range"
    p = doc.paragraphs[p_idx]
    full_text = p.text or ""
    if start_pos > len(full_text) or end_pos > len(full_text):
        return "range out of bounds"
    # Remove all existing runs
    for r in list(p.runs):
        p._element.remove(r._element)
    # Rebuild runs
    before_text = full_text[:start_pos]
    target_text = full_text[start_pos:end_pos]
    after_text = full_text[end_pos:]
    if before_text:
        p.add_run(before_text)
    target_run = p.add_run(target_text)
    fmt = target_run.font
    if bold is not None:
        target_run.bold = bool(bold)
    if italic is not None:
        target_run.italic = bool(italic)
    if underline is not None:
        target_run.underline = bool(underline)
    if color:
        try:
            fmt.color.rgb = _docx.shared.RGBColor.from_string(str(color))
        except Exception as exc:
            logging.warning(
                "Failed to set font color for %s (user %s): %s",
                blob_name,
                user_id,
                exc,
            )
    if font_size is not None:
        try:
            fmt.size = _docx.shared.Pt(float(font_size))
        except Exception as exc:
            logging.warning(
                "Failed to set font size for %s (user %s): %s",
                blob_name,
                user_id,
                exc,
            )
    if font_name:
        fmt.name = str(font_name)
    if after_text:
        p.add_run(after_text)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_delete_paragraph = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("paragraph_index", "number",
                     "Zero-based paragraph index").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_delete_paragraph",
    description="Delete a paragraph by index.",
    toolProperties=word_tool_props_delete_paragraph,
)
def word_delete_paragraph(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    p_idx = int(args.get("paragraph_index", -1))
    if not filename or p_idx < 0:
        return "Invalid arguments"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    if p_idx >= len(doc.paragraphs):
        return "paragraph_index out of range"
    p = doc.paragraphs[p_idx]
    p._element.getparent().remove(p._element)
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})


word_tool_props_create_custom_style = json.dumps(
    [
        ToolProperty("user_id", "string",
                     "User identifier used to namespace blobs").to_dict(),
        ToolProperty("filename", "string",
                     "Target .docx filename (without user prefix)").to_dict(),
        ToolProperty("style_name", "string", "New style name").to_dict(),
        ToolProperty("bold", "boolean", "Bold").to_dict(),
        ToolProperty("italic", "boolean", "Italic").to_dict(),
        ToolProperty("font_size", "number", "Font size in points").to_dict(),
        ToolProperty("font_name", "string", "Font family name").to_dict(),
        ToolProperty("color", "string", "Hex color without #").to_dict(),
        ToolProperty("base_style", "string",
                     "Base style name, e.g., Normal").to_dict(),
    ]
)


@app.generic_trigger(
    arg_name="context",
    type="mcpToolTrigger",
    toolName="word_create_custom_style",
    description="Create a custom paragraph style with basic font settings.",
    toolProperties=word_tool_props_create_custom_style,
)
def word_create_custom_style(context) -> str:
    _init_word_libs()
    _init_storage()
    payload = json.loads(context)
    args = payload.get("arguments", {})
    user_id = args.get("user_id")
    filename = args.get("filename")
    style_name = args.get("style_name")
    bold = args.get("bold")
    italic = args.get("italic")
    font_size = args.get("font_size")
    font_name = args.get("font_name")
    color = args.get("color")
    base_style = args.get("base_style") or "Normal"
    if not filename or not style_name:
        return "Missing filename or style_name"
    blob_name = f"{user_id}/{filename}" if user_id else filename
    local_path = _download_blob_to_temp(blob_name)
    doc = _docx.Document(local_path)
    try:
        # use direct import to ensure enum access
        from docx.enum.style import WD_STYLE_TYPE
        style_type = WD_STYLE_TYPE.PARAGRAPH
    except Exception:
        style_type = None
    styles = doc.styles
    try:
        style = styles[style_name]
    except Exception:
        if style_type is not None:
            style = styles.add_style(style_name, style_type)
        else:
            style = styles.add_style(style_name, 1)  # 1 likely PARAGRAPH
    try:
        base = styles[base_style]
        style.base_style = base
    except Exception as exc:
        logging.warning(
            "Failed to set base style %s for %s in file %s (user %s): %s",
            base_style,
            style_name,
            blob_name,
            user_id,
            exc,
        )
    f = style.font
    if bold is not None:
        f.bold = bool(bold)
    if italic is not None:
        f.italic = bool(italic)
    if font_size is not None:
        try:
            f.size = _docx.shared.Pt(float(font_size))
        except Exception as exc:
            logging.warning(
                "Failed to set font size for style %s in %s (user %s): %s",
                style_name,
                blob_name,
                user_id,
                exc,
            )
    if font_name:
        f.name = str(font_name)
    if color:
        try:
            f.color.rgb = _docx.shared.RGBColor.from_string(str(color))
        except Exception as exc:
            logging.warning(
                "Failed to set color for style %s in %s (user %s): %s",
                style_name,
                blob_name,
                user_id,
                exc,
            )
    doc.save(local_path)
    _upload_file_to_blob(local_path, blob_name)
    sas = _generate_blob_sas_url(blob_name, permissions="r")
    return json.dumps({"blob": blob_name, "sasUrl": sas.get("url"), "expiresUtc": sas.get("expiresUtc")})
